"""DOCX reader implementation."""

from __future__ import annotations

from pathlib import Path

from .models import DocumentContent, DocumentStats, ParagraphInfo
from .utils import count_non_whitespace_chars, detect_heading_number, looks_like_heading_text


class DocxReadError(Exception):
    """Raised when a DOCX file cannot be read."""


def _is_heading(style_name: str, text: str) -> bool:
    normalized_style = style_name.lower()
    return "heading" in normalized_style or "标题" in style_name or looks_like_heading_text(text)


def read_docx(file_path: Path) -> DocumentContent:
    """Read a .docx file and extract paragraphs and statistics."""

    try:
        from docx import Document
        from docx.opc.exceptions import PackageNotFoundError
    except ModuleNotFoundError as exc:
        raise DocxReadError("缺少依赖 python-docx，请先执行 pip install -r requirements.txt。") from exc

    try:
        document = Document(str(file_path))
    except PackageNotFoundError as exc:
        raise DocxReadError("文件不存在或不是有效的 .docx 文件。") from exc
    except Exception as exc:
        raise DocxReadError("文档读取失败，文件可能已损坏或格式异常。") from exc

    paragraphs = []
    total_chars = 0
    empty_count = 0
    heading_count = 0

    for index, paragraph in enumerate(document.paragraphs, start=1):
        text = paragraph.text or ""
        stripped = text.strip()
        style_name = paragraph.style.name if paragraph.style is not None else ""
        is_heading = _is_heading(style_name, stripped)
        heading_number = detect_heading_number(stripped) if is_heading else None

        paragraphs.append(
            ParagraphInfo(
                index=index,
                text=text,
                style_name=style_name,
                is_heading=is_heading,
                heading_number=heading_number,
            )
        )
        total_chars += count_non_whitespace_chars(text)
        if not stripped:
            empty_count += 1
        if is_heading:
            heading_count += 1

    stats = DocumentStats(
        file_name=file_path.name,
        file_path=str(file_path),
        total_chars=total_chars,
        paragraph_count=len(paragraphs),
        empty_paragraph_count=empty_count,
        heading_count=heading_count,
        table_count=len(document.tables),
    )
    return DocumentContent(stats=stats, paragraphs=paragraphs)
